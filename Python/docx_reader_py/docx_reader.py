import os
import shutil
import zipfile
import logging
from PIL import Image
from docx import Document
from datetime import datetime

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class DocxReader:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.extract_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'extract_data_py')
        self.temp_dir = os.path.join(self.extract_dir, 'temp')
        
    def extract_images(self):
        os.makedirs(self.extract_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        image_folder = os.path.join(self.extract_dir, 'images')
        os.makedirs(image_folder, exist_ok=True)
        
        image_paths = []
        
        try:
            with zipfile.ZipFile(self.docx_path) as docx:
                for file in docx.namelist():
                    if file.startswith('word/media/'):
                        image_name = os.path.basename(file)
                        docx.extract(file, self.temp_dir)
                        temp_image_path = os.path.join(self.temp_dir, 'word', 'media', image_name)
                        final_image_name = f"{timestamp}_{image_name}"
                        final_image_path = os.path.join(image_folder, final_image_name)
                        
                        try:
                            with Image.open(temp_image_path) as img:
                                if img.mode in ('RGBA', 'P'):
                                    img = img.convert('RGB')
                                img.save(final_image_path, 'JPEG', quality=85, optimize=True)
                            
                            image_paths.append({
                                'path': final_image_name,
                                'name': image_name,
                                'timestamp': timestamp
                            })
                            logging.info(f"Image saved: {final_image_path}")
                        except Exception as e:
                            logging.error(f"Error processing image {image_name}: {str(e)}")
                            continue
        
        except Exception as e:
            logging.error(f"Error extracting images: {str(e)}")
            return []
        
        finally:
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        
        return image_paths

    def extract_tables(self):
        tables = []
        try:
            doc = Document(self.docx_path)
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        row_data.append(text)
                    if any(cell for cell in row_data):
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_index': i,
                        'content': table_data,
                        'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')
                    })
            
        except Exception as e:
            logging.error(f"Error extracting tables: {str(e)}")
        
        return tables

    def extract_text(self):
        try:
            doc = Document(self.docx_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return [{'page': 1, 'text': full_text}]
            
        except Exception as e:
            logging.error(f"Error extracting text: {str(e)}")
            return []

    def extract_metadata(self):
        try:
            doc = Document(self.docx_path)
            core_properties = doc.core_properties
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            metadata = {
                'author': core_properties.author,
                'created': core_properties.created,
                'modified': core_properties.modified,
                'last_modified_by': core_properties.last_modified_by,
                'title': core_properties.title,
                'revision': core_properties.revision,
                'paragraphs_count': len(doc.paragraphs),
                'word_count': sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
                'tables_count': len(doc.tables),
                'timestamp': timestamp
            }
            
            metadata_path = os.path.join(self.extract_dir, f'metadata_{timestamp}.txt')
            os.makedirs(os.path.dirname(metadata_path), exist_ok=True)
            
            with open(metadata_path, 'w', encoding='utf-8') as f:
                for key, value in metadata.items():
                    f.write(f"{key}: {value}\n")

            metadata['saved_path'] = metadata_path
            
            return metadata
            
        except Exception as e:
            logging.error(f"Error extracting metadata: {str(e)}")
            return {}

    def process_document(self):
        try:
            images = self.extract_images()
            tables = self.extract_tables()
            texts_full = self.extract_text() 
            metadata = self.extract_metadata()
            
            return {
                'images': images,
                'tables': tables,
                'texts_full': texts_full,
                'metadata': metadata,
            }
            
        except Exception as e:
            logging.error(f"Error processing document: {str(e)}")
            return {
                'images': [], 'tables': [],
                'texts_full': [], 'metadata': {}
            }