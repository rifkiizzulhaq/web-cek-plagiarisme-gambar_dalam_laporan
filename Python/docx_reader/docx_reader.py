import os
import zipfile
from docx import Document
import shutil
from datetime import datetime
import logging
from PIL import Image

# Konfigurasi logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class DocxReader:
    def __init__(self, docx_path):
        """
        Inisialisasi DocxReader
        :param docx_path: Path ke file docx yang akan diproses
        """
        self.docx_path = docx_path
        self.extract_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'extract_data')
        self.temp_dir = os.path.join(self.extract_dir, 'temp')
        
    def extract_images(self):
        """
        Mengekstrak gambar dari file docx
        :return: List path gambar yang diekstrak
        """
        # Buat direktori jika belum ada
        os.makedirs(self.extract_dir, exist_ok=True)
        os.makedirs(self.temp_dir, exist_ok=True)
        
        # Buat subfolder dengan timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        image_folder = os.path.join(self.extract_dir, 'images')
        os.makedirs(image_folder, exist_ok=True)
        
        image_paths = []
        
        try:
            # Ekstrak file docx sebagai zip
            with zipfile.ZipFile(self.docx_path) as docx:
                # Cari semua file gambar dalam docx
                for file in docx.namelist():
                    if file.startswith('word/media/'):
                        image_name = os.path.basename(file)
                        # Ekstrak gambar ke folder temporary
                        docx.extract(file, self.temp_dir)
                        # Pindahkan ke folder final dengan nama unik
                        temp_image_path = os.path.join(self.temp_dir, 'word', 'media', image_name)
                        final_image_name = f"{timestamp}_{image_name}"
                        final_image_path = os.path.join(image_folder, final_image_name)
                        
                        # Validasi dan proses gambar
                        try:
                            with Image.open(temp_image_path) as img:
                                # Konversi ke RGB jika diperlukan
                                if img.mode in ('RGBA', 'P'):
                                    img = img.convert('RGB')
                                # Simpan dengan kualitas yang dioptimalkan
                                img.save(final_image_path, 'JPEG', quality=85, optimize=True)
                            
                            # Return only the filename for the path
                            image_paths.append({
                                'path': final_image_name,  # Just the filename
                                'name': image_name,
                                'timestamp': timestamp
                            })
                            logging.info(f"Image saved: {final_image_path}")
                        except Exception as e:
                            logging.error(f"Error processing image {image_name}: {str(e)}")
                            continue
        
        except Exception as e:
            logging.error(f"Error extracting images: {str(e)}")
            return []
        
        finally:
            # Bersihkan folder temporary
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        
        return image_paths

    def extract_tables(self):
        """
        Mengekstrak tabel dari dokumen
        :return: List tabel yang diekstrak
        """
        tables = []
        try:
            doc = Document(self.docx_path)
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Bersihkan teks dalam sel
                        text = cell.text.strip()
                        row_data.append(text)
                    if any(cell for cell in row_data):  # Skip empty rows
                        table_data.append(row_data)
                
                if table_data:  # Only add non-empty tables
                    tables.append({
                        'table_index': i,
                        'content': table_data,
                        'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')
                    })
            
        except Exception as e:
            logging.error(f"Error extracting tables: {str(e)}")
        
        return tables

    def extract_text(self):
        """
        Mengekstrak teks dari dokumen dengan mempertahankan struktur
        :return: Dictionary berisi teks terstruktur
        """
        try:
            doc = Document(self.docx_path)
            structured_text = []
            current_heading = None
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Get paragraph style
                style = paragraph.style.name
                
                # Process based on style
                if style.startswith('Heading'):
                    level = int(style.replace('Heading ', ''))
                    current_heading = {'level': level, 'text': text}
                    structured_text.append({
                        'type': 'heading',
                        'content': current_heading
                    })
                elif style.startswith('List'):
                    level = style.count('List')
                    structured_text.append({
                        'type': 'list_item',
                        'content': text,
                        'level': level
                    })
                else:
                    structured_text.append({
                        'type': 'paragraph',
                        'content': text,
                        'heading_context': current_heading['text'] if current_heading else None
                    })
            
            return {
                'content': structured_text,
                'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')
            }
            
        except Exception as e:
            logging.error(f"Error extracting text: {str(e)}")
            return {'content': [], 'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')}

    def extract_metadata(self):
        """
        Mengekstrak metadata dari file docx
        :return: Dictionary berisi metadata
        """
        try:
            doc = Document(self.docx_path)
            core_properties = doc.core_properties
            
            metadata = {
                'author': core_properties.author,
                'created': core_properties.created,
                'modified': core_properties.modified,
                'last_modified_by': core_properties.last_modified_by,
                'title': core_properties.title,
                'revision': core_properties.revision,
                'paragraphs_count': len(doc.paragraphs),
                'word_count': sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
                'tables_count': len(doc.tables),
                'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')
            }
            
            # Simpan metadata ke file
            metadata_path = os.path.join(self.extract_dir, f'metadata_{metadata["timestamp"]}.txt')
            os.makedirs(os.path.dirname(metadata_path), exist_ok=True)
            
            with open(metadata_path, 'w', encoding='utf-8') as f:
                for key, value in metadata.items():
                    f.write(f"{key}: {value}\n")
            
            return metadata
            
        except Exception as e:
            logging.error(f"Error extracting metadata: {str(e)}")
            return {}

    def process_document(self):
        """
        Memproses dokumen: mengekstrak gambar, tabel, teks, dan metadata
        :return: Dictionary berisi semua hasil ekstraksi
        """
        try:
            images = self.extract_images()
            tables = self.extract_tables()
            text = self.extract_text()
            metadata = self.extract_metadata()
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            return {
                'images': images,
                'tables': tables,
                'text': text,
                'metadata': metadata,
                'timestamp': timestamp
            }
            
        except Exception as e:
            logging.error(f"Error processing document: {str(e)}")
            return {
                'images': [],
                'tables': [],
                'text': {'content': [], 'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')},
                'metadata': {},
                'timestamp': datetime.now().strftime('%Y%m%d_%H%M%S')
            }